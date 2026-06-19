# -*- coding: utf-8 -*-
"""生成《Claude Code 与 Claude Agent SDK 能力对比》docx 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- 基础样式 ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '微软雅黑')


def _set_cn(run, font_cn='微软雅黑', font_en='Calibri'):
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), font_cn)


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _set_cn(r)
    return h


def add_p(text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    _set_cn(r)
    return p


def add_bullet(bold_prefix, text=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        _set_cn(r)
    if text:
        r2 = p.add_run(text)
        _set_cn(r2)
    return p


def add_code(code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    _set_cn(r, font_cn='Consolas', font_en='Consolas')
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_table(headers, rows, widths=None, header_fill='1F3864', header_color=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = header_color
        run.font.size = Pt(10)
        _set_cn(run)
        _shade(cell, header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            _set_cn(run)
            if ri % 2 == 1:
                _shade(cell, 'F2F2F2')
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return table


# ======================================================================
# 封面
# ======================================================================
title = doc.add_heading('Claude Code 与 Claude Agent SDK 能力对比', level=0)
for r in title.runs:
    _set_cn(r)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('交互式 CLI（harness）  vs  可编程 Agent 库（Python / TypeScript）')
rs.italic = True
rs.font.size = Pt(12)
_set_cn(rs)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run('整理日期：2026-06-15    |    信息来源：code.claude.com 官方文档')
rm.font.size = Pt(9)
rm.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
_set_cn(rm)

doc.add_paragraph()

# ======================================================================
# 一、核心结论
# ======================================================================
add_h('一、一句话核心结论', 1)
add_p('两者同源。Agent SDK 官方定位是「把驱动 Claude Code 的同一套工具、agent loop 和上下文管理，'
      '以可编程方式提供给 Python 与 TypeScript」。', bold=True)
add_p('换言之：Claude Code 是这整套能力的「交互式外壳」，Agent SDK 是同一套能力的「可编程内核」。'
      '它们不是两套独立产品，而是同一个大脑、两种身体：')
add_bullet('Claude Code = ', '交互式终端 / TUI 应用，面向人类开发者，人机对话协作。')
add_bullet('Agent SDK = ', '嵌入你自己进程的库，面向程序、服务、CI/CD，无头（headless）自动执行。')
add_p('技术注脚：TypeScript 版 SDK 会随包附带一个 native Claude Code 二进制作为可选依赖；'
      'SDK 内部本质上就是在驱动同一个 Claude Code 运行时。', italic=True, size=9.5,
      color=(0x60, 0x60, 0x60))

# ======================================================================
# 二、定位对比
# ======================================================================
add_h('二、定位与形态对比', 1)
add_table(
    ['对比维度', 'Claude Code (CLI)', 'Claude Agent SDK'],
    [
        ['产品形态', '交互式终端 / TUI 应用', '编程库（pip claude-agent-sdk / npm @anthropic-ai/claude-agent-sdk）'],
        ['面向对象', '人类开发者，对话式协作', '程序、后端服务、脚本、流水线'],
        ['运行方式', '本地终端实时交互，流式渲染 UI', '嵌入调用方进程，经 query() / ClaudeAgent 驱动'],
        ['交互模型', '人机协作：可中途介入、确认、纠偏、追问', '无头：按配置自动执行，approval 走回调'],
        ['底层运行时', '自身即「原版」', '复用 / 内嵌 Claude Code 二进制与 agent loop'],
        ['典型入口', 'claude 命令、IDE 插件、Web 版', 'query()、ClaudeAgent 类、异步流'],
        ['会话载体', '本地会话历史，按 git 项目组织', 'JSONL session 文件，可 resume / fork / 外部存储'],
        ['是否需要自己写工具循环', '否', '否（内置工具执行；区别于 Client SDK 需自写循环）'],
    ],
    widths=[Pt(120), Pt(220), Pt(240)],
)

# ======================================================================
# 三、交互与执行模型
# ======================================================================
add_h('三、交互与执行模型差异', 1)

add_h('3.1 Claude Code：人机协作的交互式 harness', 2)
add_bullet('实时 TUI：', '流式渲染思考、工具调用、diff、状态栏（statusline）。')
add_bullet('交互确认：', '敏感操作（写文件、跑命令）默认弹窗征求人类同意。')
add_bullet('Plan mode：', '先出方案待批准，再动手实现，适合复杂改动。')
add_bullet('多入口工具：', '内置 Workflow 多 agent 编排、Cron 定时任务、Worktree 隔离、后台任务、'
          'AskUserQuestion 多选澄清等交互式特性。')
add_bullet('人在回路：', '随时打断、追问、改方向；适合探索性、不确定的任务。')

add_h('3.2 Agent SDK：无头、可编排的程序化内核', 2)
add_bullet('query() 流式接口：', 'async for message 迭代每一条流式消息，结果可程序化消费。')
add_bullet('无交互 UI：', '不弹窗，靠 permission_mode（如 acceptEdits / bypassPermissions）+ allowed_tools '
          '白名单自动放行；需人类介入时用回调处理 approval。')
add_bullet('完全可控：', '可在 agent 生命周期任意点拦截、记录、阻断、改写行为。')
add_bullet('批量 / 无人值守：', '天然适配 CI/CD、定时批处理、服务端并发调度。')
add_p('最小示例（Python）：', bold=True)
add_code('''import asyncio
from claude_agent_sdk import query, ClaudeAgentOptions

async def main():
    async for message in query(
        prompt="Find and fix the bug in auth.py",
        options=ClaudeAgentOptions(allowed_tools=["Read", "Edit", "Bash"]),
    ):
        print(message)   # Claude 自主读文件、定位、改代码

asyncio.run(main())''')

# ======================================================================
# 四、能力矩阵
# ======================================================================
add_h('四、能力对比矩阵', 1)
add_p('✅ = 原生支持并典型使用；✅(强) = 该形态下的一等公民 / 体验更突出；'
      '○ = 需自行实现或仅部分支持；— = 不适用。', italic=True, size=9, color=(0x80, 0x80, 0x80))
add_table(
    ['能力', 'Claude Code', 'Agent SDK'],
    [
        ['内置工具集（Read/Write/Edit/Bash/Glob/Grep/WebSearch/WebFetch/Monitor/AskUserQuestion）', '✅', '✅'],
        ['agent loop（自主多步 + 工具循环）', '✅', '✅'],
        ['上下文管理 / 压缩 / 记忆', '✅', '✅'],
        ['MCP 外部服务连接', '✅', '✅（mcp_servers）'],
        ['Subagent（Agent 工具委派子任务）', '✅', '✅（agents 自定义定义）'],
        ['Hooks（生命周期拦截）', '✅（settings.json 配置式）', '✅(强)（回调函数式，更灵活）'],
        ['权限系统', '✅（交互弹窗确认）', '✅(强)（permission_mode + allowed_tools，无头）'],
        ['Skills / Commands / Memory(CLAUDE.md) / Plugins', '✅ 自动加载 .claude/', '✅ 默认加载，setting_sources 可控'],
        ['结构化输出（schema 强制）', '○（AskUserQuestion 多选）', '✅(强)（schema）'],
        ['流式消息消费', '✅（UI 渲染）', '✅(强)（async for 程序化处理）'],
        ['会话 resume / fork', '✅（--resume / --continue）', '✅（resume=、fork、外部存储）'],
        ['自定义工具（@tool / tool() 装饰器，进程内 MCP）', '○（需另起 MCP 进程）', '✅(强)（in-process MCP server）'],
        ['成本 / 用量追踪', 'UI 展示', '✅(强)（程序化读取）'],
        ['OpenTelemetry 可观测性', '有限', '✅(强)（一等公民）'],
        ['Workflow 多 agent 编排 / Cron / Worktree / 后台任务', '✅(强)', '○（需自行编排）'],
        ['交互式 Plan mode / 状态栏 / 实时 UI', '✅(强)', '○（需自建交互层）'],
        ['批量 / 无人值守 / CI 集成', '○（claude -p 有限）', '✅(强)（原生）'],
        ['嵌入产品 / 服务化 / 多租户', '—', '✅(强)'],
    ],
    widths=[Pt(260), Pt(150), Pt(170)],
)

# ======================================================================
# 五、扩展机制细项
# ======================================================================
add_h('五、扩展机制详解', 1)

add_h('5.1 内置工具集（两者共享）', 2)
add_table(
    ['工具', '作用'],
    [
        ['Read / Write / Edit', '读任意文件 / 建新文件 / 精确编辑已有文件'],
        ['Bash', '跑终端命令、脚本、git 操作'],
        ['Monitor', '监视后台脚本，把每行输出当作事件响应'],
        ['Glob / Grep', '按文件名模式查找 / 按正则搜内容'],
        ['WebSearch / WebFetch', '联网搜索 / 抓取并解析网页'],
        ['AskUserQuestion', '以多选项向用户提澄清问题'],
    ],
    widths=[Pt(150), Pt(360)],
)

add_h('5.2 Hooks：配置式 vs 回调式', 2)
add_p('同一套 hook 点（PreToolUse、PostToolUse、Stop、SessionStart、SessionEnd、UserPromptSubmit 等），'
      '实现方式不同：')
add_bullet('Claude Code：', '在 settings.json 里声明式配置，适合标准化、可复用的钩子。')
add_bullet('Agent SDK：', '传 Python/TS 回调函数，能直接访问上下文、做更复杂的拦截与变换。')

add_h('5.3 自定义工具：SDK 的独特优势', 2)
add_p('Agent SDK 用 @tool（Python）或 tool()（TypeScript）定义的函数，会被自动包装成一个'
      '「进程内 MCP server」——无需额外起进程，Claude 即可调用你应用内的任意函数。'
      '这是 SDK 相对 CLI 最显著的扩展性优势。')

add_h('5.4 Subagent / MCP / Skills', 2)
add_bullet('Subagent：', '两者都支持把子任务委派给带专属指令与工具集的子 agent；SDK 通过 agents 字典定义。')
add_bullet('MCP：', '两者都能连数据库、浏览器、API 等外部系统（SDK 被 2026 业界评测公认 MCP 集成最深）。')
add_bullet('Skills / Memory：', 'SDK 默认也会加载 .claude/ 下的 skills、commands、CLAUDE.md、plugins，可用 setting_sources 收窄。')

# ======================================================================
# 六、权限与安全
# ======================================================================
add_h('六、权限与安全模型', 1)
add_table(
    ['维度', 'Claude Code', 'Agent SDK'],
    [
        ['放行机制', '默认交互弹窗逐条确认；可设 allow/deny 规则', 'permission_mode + allowed_tools 白名单，无头自动放行'],
        ['典型模式', 'default（询问）/ acceptEdits / plan', 'default / acceptEdits / bypassPermissions 等'],
        ['MCP 工具权限', '按 mcp__<server>__<tool> 命名管控', '同上命名规则，可精细到单个 MCP 工具'],
        ['审批流转', '人在终端点同意/拒绝', '通过 approval 回调把决定回传'],
        ['安全边界', '由人实时把控，适合本地探索', '靠配置硬约束，适合无人值守与服务端'],
    ],
    widths=[Pt(110), Pt(220), Pt(220)],
)

# ======================================================================
# 七、会话与上下文
# ======================================================================
add_h('七、会话与上下文管理', 1)
add_bullet('Claude Code：', '本地会话历史按 git 项目根组织（如记忆挂在仓库根名下），支持 --resume/--continue。')
add_bullet('Agent SDK：', 'session 为 JSONL，可 resume 续接完整上下文、fork 探索不同分支，'
          '还可「持久化到外部存储」由你完全托管会话状态。')

# ======================================================================
# 八、可观测性
# ======================================================================
add_h('八、可观测性与运维', 1)
add_p('这是 Agent SDK 明显胜出的领域：')
add_bullet('成本 / 用量：', '可程序化读取每次调用的 token、费用，便于做预算控制与计费。')
add_bullet('OpenTelemetry：', '一等公民支持，可接入现有监控/链路追踪体系。')
add_bullet('审计：', '配合 hooks（如 PostToolUse 记录所有文件改动到 audit.log）实现合规审计。')

# ======================================================================
# 九、适用场景决策（官方表）
# ======================================================================
add_h('九、适用场景决策（官方建议）', 1)
add_p('Anthropic 官方给出的选型对照：', bold=True)
add_table(
    ['使用场景', '推荐选择'],
    [
        ['交互式日常开发', 'Claude Code (CLI)'],
        ['CI / CD 流水线', 'Agent SDK'],
        ['自定义应用 / 产品', 'Agent SDK'],
        ['一次性任务', 'Claude Code (CLI)'],
        ['生产自动化', 'Agent SDK'],
    ],
    widths=[Pt(260), Pt(250)],
)
add_p('官方建议：很多团队两者都用——日常开发用 CLI，生产用 SDK，工作流在两者间可直接迁移。', italic=True)

# ======================================================================
# 十、计费与认证注意事项
# ======================================================================
add_h('十、计费与认证注意事项', 1)
add_bullet('计费分离：', '自 2026-06-15 起，订阅计划下的 Agent SDK 与 claude -p 用量，'
          '从「独立的月度 Agent SDK 额度」中扣除，与交互式用量分开计量。')
add_bullet('认证方式：', 'SDK 用 ANTHROPIC_API_KEY；亦支持 Bedrock / Claude Platform on AWS / Vertex AI / Azure Foundry。')
add_bullet('第三方限制：', '未经 Anthropic 批准，第三方不得在基于 Agent SDK 的产品里提供 claude.ai 登录或速率限制。')
add_bullet('环境耦合（本项目）：', '本环境走内网代理（ANTHROPIC_BASE_URL 指向内网、后端 glm-5.2、无 sk-ant- key），'
          'SDK 的认证与端点需匹配该代理配置，不能直接套用官方 API key 流程。')

# ======================================================================
# 十一、与第三种形态的对照（延伸）
# ======================================================================
add_h('十一、延伸：还有第三种——Managed Agents', 1)
add_p('除 CLI 与 SDK 外，Anthropic 还提供 Managed Agents（托管 REST API）。三者的层级关系：')
add_table(
    ['', 'Agent SDK', 'Managed Agents'],
    [
        ['运行位置', '你自己的进程 / 基础设施', 'Anthropic 托管的基础设施'],
        ['接口', 'Python / TypeScript 库', 'REST API'],
        ['作用对象', '你本机的文件与服务', '每个 session 一个托管沙箱'],
        ['会话状态', '本地文件系统 JSONL', 'Anthropic 托管的事件日志'],
        ['自定义工具', '进程内函数（in-process）', 'Claude 触发、你执行并回传结果'],
        ['最佳用途', '本地原型、直接操作本机文件与服务', '无需自建沙箱/会话基础设施的生产 agent'],
    ],
    widths=[Pt(100), Pt(260), Pt(260)],
)
add_p('常见路径：用 Agent SDK 本地原型验证 → 生产迁到 Managed Agents。', italic=True)

# ======================================================================
# 十二、选型总结
# ======================================================================
add_h('十二、选型总结', 1)
add_bullet('选 Claude Code：', '你要的是「和人一起干活」——探索、调试、复杂重构、需要随时纠偏、'
          '想用 Workflow/Plan/状态栏这些交互式特性。')
add_bullet('选 Agent SDK：', '你要的是「让程序自己干活」——嵌入产品、跑 CI、批量任务、'
          '需要自定义工具、结构化输出、成本/可观测性控制、服务端多租户。')
add_bullet('两者并用：', '日常 CLI + 生产 SDK，是官方推荐的主流姿势；底层能力一致，迁移成本低。')

add_p('')
note = doc.add_paragraph()
rn = note.add_run('附注：本对比基于 code.claude.com 官方文档整理；具体 API 形态（query / ClaudeAgentOptions / '
                  'ClaudeAgent 等）随版本演进，使用时以最新 SDK reference 为准。')
rn.italic = True
rn.font.size = Pt(9)
rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
_set_cn(rn)

# ---------- 保存 ----------
out = r'D:\我的个人区间\Projects\context-lab\backend\sandbox\Claude Code 与 Agent SDK 能力对比.docx'
doc.save(out)
print('SAVED:', out)
