"""生成《冒泡排序逐行讲解》Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CODE_FONT = "Consolas"   # 代码 / 等宽字体
BODY_FONT = "微软雅黑"    # 正文中文字体


def set_run_font(run, font=BODY_FONT, size=11, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 确保中文字体也生效
    from docx.oxml.ns import qn
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_code(doc, text):
    """添加一段等宽字体的代码段落，带浅色块效果。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font=CODE_FONT, size=10, color=(30, 30, 30))
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_run_font(run, font=BODY_FONT, size=16 if level == 1 else 13, bold=True,
                 color=(0x1F, 0x3A, 0x68))
    return h


def add_explain(doc, line_no, code, explain):
    """添加一条逐行讲解：行号 + 代码 + 说明。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"第 {line_no} 行　")
    set_run_font(r1, font=BODY_FONT, size=11, bold=True, color=(0x00, 0x66, 0xCC))
    r2 = p.add_run(code)
    set_run_font(r2, font=CODE_FONT, size=10, color=(40, 40, 40))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Pt(18)
    r3 = p2.add_run(explain)
    set_run_font(r3, font=BODY_FONT, size=11)


def main():
    doc = Document()

    # ===== 封面标题 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("冒泡排序（Bubble Sort）逐行讲解")
    set_run_font(trun, font=BODY_FONT, size=22, bold=True, color=(0x1F, 0x3A, 0x68))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("基于 bubble_sort.py 源码的逐行解析")
    set_run_font(srun, font=BODY_FONT, size=12, color=(120, 120, 120))

    doc.add_paragraph()  # 空行

    # ===== 一、算法简介 =====
    add_heading(doc, "一、算法简介", level=1)
    intro = doc.add_paragraph()
    irun = intro.add_run(
        "冒泡排序是一种简单的比较类排序算法。其核心思想是：每轮从头到尾两两比较相邻元素，"
        "若顺序错误就交换，使每轮都将当前未排序部分的最大值「冒泡」到末尾。重复若干轮后整个列表有序。"
        "配合 swapped 标志可在已经有序时提前退出，提升最好情况下的效率。"
    )
    set_run_font(irun, font=BODY_FONT, size=11)

    # ===== 二、完整源代码 =====
    add_heading(doc, "二、完整源代码", level=1)
    source = '''"""冒泡排序（Bubble Sort）示例程序"""


def bubble_sort(arr: list) -> list:
    """对列表进行升序冒泡排序，返回排序后的新列表。

    时间复杂度：O(n^2)
    空间复杂度：O(1)（原地排序，这里为不修改原列表做了拷贝）
    """
    # 拷贝一份，避免修改原始数据
    nums = arr[:]
    n = len(nums)

    for i in range(n - 1):
        swapped = False  # 优化：若本轮没有发生交换，说明已有序，可提前结束
        for j in range(n - 1 - i):
            if nums[j] > nums[j + 1]:
                nums[j], nums[j + 1] = nums[j + 1], nums[j]
                swapped = True
        if not swapped:
            break

    return nums


def main():
    data = [64, 34, 25, 12, 22, 11, 90]
    print("排序前：", data)
    print("排序后：", bubble_sort(data))
    # 原列表不受影响
    print("原列表：", data)


if __name__ == "__main__":
    main()
'''
    add_code(doc, source)

    # ===== 三、逐行讲解 =====
    add_heading(doc, "三、逐行讲解", level=1)

    add_explain(doc, 1, '"""冒泡排序（Bubble Sort）示例程序"""',
                "模块级文档字符串（docstring），用三引号包裹，说明本文件的用途。")

    add_explain(doc, 4, "def bubble_sort(arr: list) -> list:",
                "定义函数 bubble_sort，接收列表参数 arr。arr: list 表示参数类型为 list，"
                "-> list 表示返回值也是 list。这些是类型注解，仅供提示，不影响运行。")

    add_explain(doc, "5–9", "函数的文档字符串",
                "说明函数用途、时间复杂度 O(n²)、空间复杂度 O(1)。")

    add_explain(doc, 10, "# 拷贝一份，避免修改原始数据",
                "注释，解释下方 arr[:] 的目的。")

    add_explain(doc, 11, "nums = arr[:]",
                "arr[:] 是切片语法，表示取整个列表的副本，赋值给 nums。后续排序只改 nums，"
                "不会影响原始数据。")

    add_explain(doc, 12, "n = len(nums)",
                "取列表长度存到变量 n。")

    add_explain(doc, 14, "for i in range(n - 1):",
                "外层循环，控制轮数。n 个元素最多需要 n−1 轮比较才能确保排好序。")

    add_explain(doc, 15, "swapped = False",
                "每轮开始前，把「本轮是否发生过交换」标志设为 False。这是个优化标志。")

    add_explain(doc, 16, "for j in range(n - 1 - i):",
                "内层循环，进行相邻元素两两比较。每完成一轮，最大元素已冒泡到末尾，"
                "所以下一轮可少比一个，用 n - 1 - i 实现。")

    add_explain(doc, 17, "if nums[j] > nums[j + 1]:",
                "如果前一个元素比后一个大（不符合升序）……")

    add_explain(doc, 18, "nums[j], nums[j + 1] = nums[j + 1], nums[j]",
                "Python 的多重赋值，交换这两个元素的位置（让小的靠前）。")

    add_explain(doc, 19, "swapped = True",
                "记录本轮发生过交换。")

    add_explain(doc, "20–21", "if not swapped: break",
                "如果这一轮一次交换都没发生，说明列表已完全有序，直接 break 跳出外层循环，"
                "提前结束，避免无意义的后续轮次。")

    add_explain(doc, 23, "return nums",
                "返回排好序的新列表。")

    add_explain(doc, 26, "def main():",
                "定义主函数，用来演示调用。")

    add_explain(doc, 27, 'data = [64, 34, 25, 12, 22, 11, 90]',
                "准备一组测试数据。")

    add_explain(doc, 28, 'print("排序前：", data)',
                "打印排序前的原始列表。")

    add_explain(doc, 29, 'print("排序后：", bubble_sort(data))',
                "调用 bubble_sort 并打印排序结果。")

    add_explain(doc, "30–31", 'print("原列表：", data)',
                "注释说明原列表不会被改动，并打印原列表验证这一点。")

    add_explain(doc, "34–35", 'if __name__ == "__main__": main()',
                "Python 的惯用入口写法：只有直接运行此文件时才执行 main()；"
                "被 import 引入时不会自动运行。")

    # ===== 四、核心思路总结 =====
    add_heading(doc, "四、核心思路总结", level=1)
    points = [
        "每轮从头到尾两两比较相邻元素，把较大的往后换。",
        "每轮都会把当前未排序部分的最大值「冒泡」到末尾。",
        "配合 swapped 标志，在列表已经有序时可提前退出，减少无谓比较。",
        "时间复杂度 O(n²)，空间复杂度 O(1)；最好情况（已有序）为 O(n)。",
    ]
    for pt in points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(pt)
        set_run_font(run, font=BODY_FONT, size=11)

    # ===== 保存 =====
    out = "bubble_sort_讲解.docx"
    doc.save(out)
    print("已生成：", out)


if __name__ == "__main__":
    main()
